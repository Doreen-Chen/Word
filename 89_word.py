# 89. 處理 Word 檔 (使用第三方套件流程示範)
#   https://python-docx.readthedocs.io/en/latest/

# from 指定套件 ; import class
from docx import Document
from docx.shared import Inches

# new object
document = Document()

# Word : Heading
document.add_heading('Document Title', 0)

# Word : Paragraph
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# Word : Heading
document.add_heading('Heading, level 1', level=1)

# Word : Paragraph
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# Word : Picture
document.add_picture('89_test.png', width=Inches(1.25))  # 這裡要換圖片檔名

# Data array
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# 新增 Table
table = document.add_table(rows=1, cols=3)  # 先新增第一列欄位名稱
hdr_cells = table.rows[0].cells  # first row is Column Header
hdr_cells[0].text = 'Qty'  # Column Name 1st
hdr_cells[1].text = 'Id'  # Column Name 2nd
hdr_cells[2].text = 'Desc'  # Column Name 3rd
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')